"""
文档处理器 - 处理文档格式转换
专门负责JSON数据与DOCX格式之间的转换

调用关系:
- core/pipeline.py: 调用export_script_to_docx导出阅读版DOCX文档
- 可被用户通过CLI或Web界面间接调用来进行文档格式转换
- 从utils.py迁移而来，专注于DOCX文档的导出和解析功能
"""

import os
import re
from typing import Dict, Any, List

from core.domain.metadata import (
    get_source_name,
    get_video_titles,
    get_cover_titles,
    get_cover_subtitles,
    get_golden_quotes,
    get_primary_video_title,
)
from core.shared import logger, ensure_directory_exists, FileProcessingError


OPTION_LABELS = {
    "video_title": {
        "label": "VIDEO_TITLE",
        "display": "视频标题",
    },
    "cover_title": {
        "label": "COVER_TITLE",
        "display": "封面主标题",
    },
    "cover_subtitle": {
        "label": "COVER_SUBTITLE",
        "display": "封面副标题",
    },
    "golden_quote": {
        "label": "GOLDEN_QUOTE",
        "display": "开场金句",
    },
}
OPTION_HEADER_TEMPLATE = ">>> {label} OPTION {index:02d} >>>"
OPTION_MARKERS = {
    "video_title": ("===VIDEO_TITLES_START===", "===VIDEO_TITLES_END==="),
    "cover_title": ("===COVER_TITLES_START===", "===COVER_TITLES_END==="),
    "cover_subtitle": ("===COVER_SUBTITLES_START===", "===COVER_SUBTITLES_END==="),
    "golden_quote": ("===GOLDEN_QUOTES_START===", "===GOLDEN_QUOTES_END==="),
}
OPTION_VALUE_KEYS = {
    "video_title": ("video_titles",),
    "cover_title": ("cover_titles",),
    "cover_subtitle": ("cover_subtitles",),
    "golden_quote": ("golden_quotes",),
}


def _dedupe_options(values: List[str]) -> List[str]:
    result: List[str] = []
    for value in values or []:
        if isinstance(value, str):
            candidate = value.strip()
            if candidate and candidate not in result:
                result.append(candidate)
    return result


def _prepare_option_values(raw_data: Dict[str, Any], field: str) -> List[str]:
    prepared: List[str] = []
    for key in OPTION_VALUE_KEYS.get(field, (f"{field}_options", field)):
        value = raw_data.get(key)
        if isinstance(value, list):
            for item in value:
                if isinstance(item, str):
                    candidate = item.strip()
                    if candidate and candidate not in prepared:
                        prepared.append(candidate)
        elif isinstance(value, str):
            candidate = value.strip()
            if candidate and candidate not in prepared:
                prepared.append(candidate)
    return prepared if prepared else [""]


def _write_option_block(document, field: str, options: List[str]):
    meta = OPTION_LABELS[field]
    header_label = meta["label"]
    display_label = meta["display"]
    instruction = document.add_paragraph(f"请保留以下 '>>> {header_label} OPTION XX >>>' 作为{display_label}的分隔符，可调整文本或顺序；系统默认使用第一条作为主用内容。")
    _setup_docx_paragraph(instruction)

    for idx, option in enumerate(options, 1):
        header_text = OPTION_HEADER_TEMPLATE.format(label=header_label, index=idx)
        header_para = document.add_paragraph(header_text)
        _setup_docx_paragraph(header_para)

        value_para = document.add_paragraph(option)
        _setup_docx_paragraph(value_para)


def _extract_option_values(paragraphs: List[str], start_idx: int, end_idx: int, field: str) -> List[str]:
    if start_idx == -1 or end_idx == -1 or end_idx <= start_idx + 1:
        return []

    label = OPTION_LABELS[field]["label"]
    pattern = re.compile(rf"^>>> {label} OPTION \d{{1,2}} >>>$")

    collected: List[str] = []
    buffer: List[str] = []
    collecting = False

    for text in paragraphs[start_idx + 1:end_idx]:
        if pattern.match(text):
            if collecting:
                option_text = "\n".join(buffer).strip()
                if option_text:
                    collected.append(option_text)
                buffer = []
            collecting = True
            continue

        if collecting:
            buffer.append(text)

    if collecting:
        option_text = "\n".join(buffer).strip()
        if option_text:
            collected.append(option_text)

    if not collected:
        raw_block = [line for line in paragraphs[start_idx + 1:end_idx] if not line.startswith("请保留以下 ")]
        fallback_text = "\n".join(raw_block).strip()
        if fallback_text:
            collected.append(fallback_text)

    return _dedupe_options(collected)


def export_script_to_docx(script_data: Dict[str, Any], docx_path: str) -> str:
    """
    将脚本JSON导出为可阅读的DOCX文档
    
    Args:
        script_data: 含有 title 与 segments 的脚本数据
        docx_path: 输出的docx文件完整路径
        
    Returns:
        str: 实际保存的docx路径
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    _setup_docx_style(document)

    title_text = get_primary_video_title(script_data, 'untitled')
    segments = script_data.get('segments', []) or []

    # 标题（居中）
    title_para = document.add_paragraph()
    title_run = title_para.add_run(title_text)
    _setup_docx_run(title_run)
    _setup_docx_paragraph(title_para, WD_ALIGN_PARAGRAPH.CENTER)

    # 正文段落（两端对齐）
    for seg in segments:
        content = (seg or {}).get('content', '')
        if not content:
            continue
        p = document.add_paragraph()
        r = p.add_run(content)
        _setup_docx_run(r)
        _setup_docx_paragraph(p, WD_ALIGN_PARAGRAPH.JUSTIFY)

    ensure_directory_exists(os.path.dirname(docx_path))
    document.save(docx_path)
    logger.info(f"阅读版DOCX已保存: {docx_path}")
    return docx_path


def export_raw_to_docx(raw_data: Dict[str, Any], docx_path: str) -> str:
    """
    将原始LLM输出的JSON数据导出为带标记的DOCX文档，方便用户编辑
    
    Args:
        raw_data: 包含title、golden_quote、content的原始数据
        docx_path: 输出的docx文件路径
        
    Returns:
        str: 实际保存的docx路径
    """
    from docx import Document

    document = Document()
    _setup_docx_style(document)

    # 添加编辑说明
    instruction_para = document.add_paragraph()
    instruction_run = instruction_para.add_run("编辑说明：请直接编辑下方内容，但请保持标记符号（===...===）不变，这些标记用于系统识别各个字段。")
    _setup_docx_run(instruction_run)
    _setup_docx_paragraph(instruction_para)
    
    source_name = get_source_name(raw_data)

    # 原始作品标题
    for start_marker, content, end_marker in [
        ("===SOURCE_NAME_START===", source_name, "===SOURCE_NAME_END==="),
    ]:
        document.add_paragraph()

        start_para = document.add_paragraph(start_marker)
        _setup_docx_paragraph(start_para)

        content_para = document.add_paragraph(content)
        _setup_docx_paragraph(content_para)

        end_para = document.add_paragraph(end_marker)
        _setup_docx_paragraph(end_para)

    # 多选字段：视频标题、封面主标题、封面副标题与开场金句
    for field in ("video_title", "cover_title", "cover_subtitle", "golden_quote"):
        document.add_paragraph()

        start_marker, end_marker = OPTION_MARKERS[field]

        start_para = document.add_paragraph(start_marker)
        _setup_docx_paragraph(start_para)

        options = _prepare_option_values(raw_data, field)
        _write_option_block(document, field, options)

        end_para = document.add_paragraph(end_marker)
        _setup_docx_paragraph(end_para)

    # 正文内容
    document.add_paragraph()

    content_start = document.add_paragraph("===CONTENT_START===")
    _setup_docx_paragraph(content_start)

    content_para = document.add_paragraph(raw_data.get('content', ''))
    _setup_docx_paragraph(content_para)

    content_end = document.add_paragraph("===CONTENT_END===")
    _setup_docx_paragraph(content_end)

    # 统一设置所有段落的字体
    for para in document.paragraphs:
        for run in para.runs:
            _setup_docx_run(run)

    ensure_directory_exists(os.path.dirname(docx_path))
    document.save(docx_path)
    logger.info(f"原始编辑版DOCX已保存: {docx_path}")
    return docx_path


def parse_raw_from_docx(docx_path: str) -> Dict[str, Any]:
    """
    从带标记的DOCX文档解析回原始数据格式
    
    Args:
        docx_path: 包含标记的docx文件路径
        
    Returns:
        Dict[str, Any]: 解析得到的原始数据，包含title、golden_quote、content
        
    Raises:
        FileProcessingError: 解析失败时抛出
    """
    from docx import Document
    
    if not os.path.exists(docx_path):
        raise FileProcessingError(f"DOCX文件不存在: {docx_path}")
    
    try:
        document = Document(docx_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:  # 只保留非空段落
                paragraphs.append(text)
        
        # 查找标记位置
        source_name_start = source_name_end = video_titles_start = video_titles_end = -1
        cover_titles_start = cover_titles_end = cover_subtitles_start = cover_subtitles_end = -1
        golden_quotes_start = golden_quotes_end = content_start = content_end = -1

        for i, para_text in enumerate(paragraphs):
            if para_text == "===SOURCE_NAME_START===":
                source_name_start = i
            elif para_text == "===SOURCE_NAME_END===":
                source_name_end = i
            elif para_text == "===VIDEO_TITLES_START===":
                video_titles_start = i
            elif para_text == "===VIDEO_TITLES_END===":
                video_titles_end = i
            elif para_text == "===COVER_TITLES_START===":
                cover_titles_start = i
            elif para_text == "===COVER_TITLES_END===":
                cover_titles_end = i
            elif para_text == "===COVER_SUBTITLES_START===":
                cover_subtitles_start = i
            elif para_text == "===COVER_SUBTITLES_END===":
                cover_subtitles_end = i
            elif para_text == "===GOLDEN_QUOTES_START===":
                golden_quotes_start = i
            elif para_text == "===GOLDEN_QUOTES_END===":
                golden_quotes_end = i
            elif para_text == "===CONTENT_START===":
                content_start = i
            elif para_text == "===CONTENT_END===":
                content_end = i

        # 验证标记完整性
        if source_name_start == -1 or source_name_end == -1:
            raise ValueError("缺少SOURCE_NAME标记")
        if video_titles_start == -1 or video_titles_end == -1:
            raise ValueError("缺少VIDEO_TITLES标记")
        if cover_titles_start == -1 or cover_titles_end == -1:
            raise ValueError("缺少COVER_TITLES标记")
        if cover_subtitles_start == -1 or cover_subtitles_end == -1:
            raise ValueError("缺少COVER_SUBTITLES标记")
        if golden_quotes_start == -1 or golden_quotes_end == -1:
            raise ValueError("缺少GOLDEN_QUOTES标记")
        if content_start == -1 or content_end == -1:
            raise ValueError("缺少CONTENT标记")

        # 提取各字段内容
        content = '\n'.join(paragraphs[content_start + 1:content_end]).strip()

        video_titles = _extract_option_values(paragraphs, video_titles_start, video_titles_end, "video_title")
        source_name = '\n'.join(paragraphs[source_name_start + 1:source_name_end]).strip()
        cover_titles = _extract_option_values(paragraphs, cover_titles_start, cover_titles_end, "cover_title")
        cover_subtitles = _extract_option_values(paragraphs, cover_subtitles_start, cover_subtitles_end, "cover_subtitle")
        golden_quotes = _extract_option_values(paragraphs, golden_quotes_start, golden_quotes_end, "golden_quote")

        result = {
            'source_name': source_name,
            'video_titles': video_titles,
            'cover_titles': cover_titles,
            'cover_subtitles': cover_subtitles,
            'golden_quotes': golden_quotes,
            'content': content
        }
        
        logger.info(f"从DOCX解析原始数据成功: {docx_path}")
        return result
        
    except Exception as e:
        logger.error(f"解析DOCX文件失败: {str(e)}")
        raise FileProcessingError(f"解析DOCX文件失败: {str(e)}")


def _setup_docx_style(document, style_name: str = 'Normal'):
    """设置DOCX文档样式"""
    try:
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        
        style = document.styles[style_name]
        style.font.name = '宋体'
        
        # 设置东亚字体
        if hasattr(style, 'element') and style.element is not None:
            rPr = style.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 行距 1.5 倍
        if style.paragraph_format is not None:
            try:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            except Exception:
                style.paragraph_format.line_spacing = 1.5
    except Exception:
        pass  # 样式设置失败也不影响功能


def _setup_docx_run(run):
    """设置DOCX文本run的字体"""
    try:
        from docx.oxml.ns import qn
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass


def _setup_docx_paragraph(para, alignment=None):
    """设置DOCX段落的格式"""
    try:
        from docx.enum.text import WD_LINE_SPACING
        if alignment:
            para.alignment = alignment
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except Exception:
        try:
            para.paragraph_format.line_spacing = 1.5
        except Exception:
            pass
