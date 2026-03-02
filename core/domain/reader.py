"""
文档读取器 - 统一的文档解析接口（迁移至 core）
支持 PDF、EPUB、MOBI、AZW3、DOCX、DOC 等格式
"""

import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Tuple

# 文档处理库
import ebooklib
from ebooklib import epub
import pdfplumber
from docx import Document

# 可选依赖
try:
    import fitz  # pymupdf
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

from core.shared import logger, FileProcessingError


def clean_text(text: str) -> str:
    """清理PDF/EPUB解析产生的乱码文本"""
    if not text:
        return ""
    
    # 移除HTML标签
    text = re.sub(r'<[^>]+>', '', text)
    
    # 清理PDF CID字符乱码：移除 (cid:数字) 格式的字符
    text = re.sub(r'\(cid:\d+\)', '', text)
    
    # 清理其他常见的PDF解析问题
    # 移除单独的数字和字母组合（可能是字体编码残留）
    text = re.sub(r'\b[A-Z]{1,3}\d*\b', ' ', text)
    
    # 更强力的乱码字符清理
    # 移除明显的非文本字符（保留中文、英文、数字、常见标点）
    def is_valid_char(char):
        # 中文字符
        if '\u4e00' <= char <= '\u9fff':
            return True
        # 英文字母和数字
        if char.isalnum() and ord(char) < 128:
            return True
        # 常见标点符号
        if char in '，。！？；：""''（）【】《》、—…·.,:;!?()[]{}"-\'':
            return True
        # 空格和换行
        if char in ' \n\t\r':
            return True
        return False
    
    # 字符级过滤
    filtered_chars = []
    for char in text:
        if is_valid_char(char):
            filtered_chars.append(char)
        else:
            # 用空格替换无效字符
            if filtered_chars and filtered_chars[-1] != ' ':
                filtered_chars.append(' ')
    
    text = ''.join(filtered_chars)
    
    # 标准化空白字符
    text = re.sub(r'\s+', ' ', text)
    # 移除首尾空白
    text = text.strip()
    
    return text


class DocumentReader:
    """统一的文档读取器"""
    
    # 支持的文件格式
    SUPPORTED_FORMATS = ['.epub', '.pdf', '.mobi', '.azw3', '.docx', '.doc']
    
    def __init__(self):
        """初始化文档读取器"""
        pass
    
    def read(self, file_path: str) -> Tuple[str, int]:
        """
        读取文档内容
        
        Args:
            file_path: 文档路径
            
        Returns:
            Tuple[str, int]: (清理后的文本内容, 字符数)
            
        Raises:
            FileProcessingError: 文件读取或处理失败
        """
        # 基本验证
        self._validate_file(file_path)
        
        # 根据文件扩展名选择读取方法
        ext = Path(file_path).suffix.lower()
        
        try:
            if ext == '.epub':
                return self._read_epub(file_path)
            elif ext == '.pdf':
                return self._read_pdf(file_path)
            elif ext == '.mobi':
                return self._read_mobi(file_path)
            elif ext == '.azw3':
                return self._read_azw3(file_path)
            elif ext == '.docx':
                return self._read_docx(file_path)
            elif ext == '.doc':
                return self._read_doc(file_path)
            else:
                raise FileProcessingError(f"不支持的文件格式: {ext}")
                
        except FileProcessingError:
            raise
        except Exception as e:
            raise FileProcessingError(f"读取文件失败: {str(e)}")
    
    def _validate_file(self, file_path: str) -> None:
        """验证文件有效性"""
        if not file_path or not file_path.strip():
            raise FileProcessingError("文件路径不能为空")
        
        if not os.path.exists(file_path):
            raise FileProcessingError(f"文件不存在: {file_path}")
        
        if not os.path.isfile(file_path):
            raise FileProcessingError(f"路径不是有效文件: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise FileProcessingError(f"不支持的文件格式: {ext}，支持的格式: {self.SUPPORTED_FORMATS}")
    
    def _read_epub(self, file_path: str) -> Tuple[str, int]:
        """读取EPUB文件"""
        logger.info(f"读取EPUB文件: {os.path.basename(file_path)}")
        
        try:
            book = epub.read_epub(file_path)
            content_parts = []
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    content = item.get_content().decode('utf-8', errors='ignore')
                    cleaned = clean_text(content)
                    if cleaned:
                        content_parts.append(cleaned)
            
            if not content_parts:
                raise FileProcessingError("EPUB文件中未找到可读取的文本内容")
            
            full_content = ' '.join(content_parts)
            word_count = len(full_content)
            
            logger.info(f"EPUB读取成功，字数: {word_count:,}")
            return full_content, word_count
            
        except Exception as e:
            raise FileProcessingError(f"EPUB读取失败: {str(e)}")
    
    def _read_pdf(self, file_path: str) -> Tuple[str, int]:
        """读取PDF文件，优先使用PyMuPDF，备用pdfplumber"""
        logger.info(f"读取PDF文件: {os.path.basename(file_path)}")
        
        content_parts = []
        
        # 优先使用PyMuPDF
        if HAS_FITZ:
            try:
                doc = fitz.open(file_path)
                for i in range(len(doc)):
                    page = doc.load_page(i)
                    text = page.get_text()
                    if text.strip():
                        content_parts.append(text)
                doc.close()
            except Exception:
                logger.debug("PyMuPDF读取失败，尝试pdfplumber")
                content_parts = []
        
        # 备用pdfplumber
        if not content_parts:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
            except Exception:
                logger.debug("pdfplumber读取失败")
        
        if not content_parts:
            raise FileProcessingError("PDF文件无法提取文本，可能是扫描版或编码问题")
        
        # 文本清理和质量检查
        full_content = ' '.join(content_parts)
        cleaned_content = clean_text(full_content)
        
        self._validate_pdf_quality(full_content, cleaned_content)
        
        logger.info(f"PDF读取成功，字数: {len(cleaned_content):,}")
        return cleaned_content, len(cleaned_content)
    
    def _validate_pdf_quality(self, original: str, cleaned: str) -> None:
        """验证PDF文本质量"""
        total_chars = len(original)
        cleaned_chars = len(cleaned)
        
        if total_chars == 0:
            raise FileProcessingError("PDF文件内容为空")
        
        # 计算损失率和可读性
        loss_ratio = (total_chars - cleaned_chars) / total_chars
        chinese_chars = sum(1 for c in cleaned if '\u4e00' <= c <= '\u9fff')
        english_chars = sum(1 for c in cleaned if c.isalpha() and ord(c) < 128)
        readable_chars = chinese_chars + english_chars
        readable_ratio = readable_chars / max(1, cleaned_chars)
        
        # 严格质量检查
        if (loss_ratio > 0.8 or readable_ratio < 0.5 or 
            cleaned_chars < 1000 or readable_chars < 5000):
            raise FileProcessingError(
                "PDF文本质量不足，可能为扫描版或编码异常。建议使用OCR识别或转换为其他格式"
            )
    
    def _read_mobi(self, file_path: str) -> Tuple[str, int]:
        """读取MOBI文件（简化实现）"""
        logger.info(f"读取MOBI文件: {os.path.basename(file_path)}")
        
        try:
            with open(file_path, 'rb') as f:
                # 检查MOBI标识
                f.seek(60)
                if f.read(4) != b'MOBI':
                    # 检查BOOKMOBI标识
                    f.seek(0)
                    content = f.read(100)
                    if b'BOOKMOBI' not in content:
                        raise FileProcessingError("不是有效的MOBI格式")
                
                # 简单文本提取
                f.seek(0)
                raw_content = f.read()
                
                # 提取可读文本
                text_content = self._extract_mobi_text(raw_content)
                
                if not text_content.strip():
                    raise FileProcessingError("MOBI文件无法提取文本内容")
                
                cleaned_content = clean_text(text_content)
                word_count = len(cleaned_content)
                
                if word_count < 100:
                    raise FileProcessingError("MOBI文件内容过少，可能是加密文件")
                
                logger.info(f"MOBI读取成功，字数: {word_count:,}")
                return cleaned_content, word_count
                
        except Exception as e:
            raise FileProcessingError(f"MOBI读取失败: {str(e)}。建议转换为EPUB格式")
    
    def _extract_mobi_text(self, raw_content: bytes) -> str:
        """从MOBI原始内容中提取文本"""
        logger.info(f"开始提取MOBI文本，文件大小: {len(raw_content)} bytes")
        
        try:
            # 尝试使用mobi库解析
            import mobi
            import tempfile
            import os
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.mobi', delete=False) as temp_file:
                temp_file.write(raw_content)
                temp_file_path = temp_file.name
            
            try:
                # 使用mobi库提取
                tempdir, filepath = mobi.extract(temp_file_path)
                
                # 读取提取的HTML文件
                text_content = ""
                for root, dirs, files in os.walk(tempdir):
                    for file in files:
                        if file.endswith(('.html', '.htm')):
                            file_path = os.path.join(root, file)
                            try:
                                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                    html_content = f.read()
                                    # 简单去除HTML标签
                                    import re
                                    clean_content = re.sub(r'<[^>]+>', ' ', html_content)
                                    clean_content = re.sub(r'\s+', ' ', clean_content)
                                    text_content += clean_content + " "
                            except:
                                continue
                
                # 清理临时文件
                import shutil
                shutil.rmtree(tempdir, ignore_errors=True)
                os.unlink(temp_file_path)
                
                if text_content.strip():
                    logger.info(f"使用mobi库提取完成，文本长度: {len(text_content)} 字符")
                    return text_content.strip()
                    
            except Exception as e:
                logger.warning(f"mobi库提取失败: {e}")
                os.unlink(temp_file_path)
                
        except ImportError:
            logger.info("mobi库未安装，使用备用方法")
        except Exception as e:
            logger.warning(f"mobi库方法失败: {e}")
        
        # 备用方法：改进的分块提取
        logger.info("使用备用文本提取方法")
        text_parts = []
        
        # 寻找可能的文本区域
        content_str = raw_content.decode('utf-8', errors='ignore')
        
        # 查找段落模式
        paragraphs = []
        lines = content_str.split('\n')
        
        for line in lines:
            line = line.strip()
            # 过滤明显的非文本内容
            if len(line) < 10 or len(line) > 1000:
                continue
            # 检查是否包含足够的字母
            alpha_count = sum(1 for c in line if c.isalpha())
            if alpha_count < len(line) * 0.5:
                continue
            # 检查是否包含正常的句子结构
            if any(word in line.lower() for word in ['the', 'and', 'that', 'this', 'with', 'for']) or \
               any(char in line for char in ['。', '，', '、']):
                paragraphs.append(line)
        
        # 限制输出长度，避免重复内容
        unique_paragraphs = []
        seen = set()
        
        for para in paragraphs:
            # 简单去重
            para_key = para[:50] if len(para) > 50 else para
            if para_key not in seen:
                seen.add(para_key)
                unique_paragraphs.append(para)
                
            # 限制总长度
            if len(' '.join(unique_paragraphs)) > 500000:  # 限制50万字符
                break
        
        result = ' '.join(unique_paragraphs)
        logger.info(f"备用方法提取完成，文本长度: {len(result)} 字符")
        return result
    
    def _read_azw3(self, file_path: str) -> Tuple[str, int]:
        """读取AZW3文件，优先使用mobi库"""
        logger.info(f"读取AZW3文件: {os.path.basename(file_path)}")
        
        try:
            # 尝试使用mobi库
            import mobi
            import shutil
            
            tempdir, filepath = mobi.extract(file_path)
            
            try:
                ext = os.path.splitext(filepath)[1].lower()
                
                if ext == '.epub':
                    content, word_count = self._read_epub(filepath)
                elif ext in ['.html', '.htm']:
                    content, word_count = self._read_html(filepath)
                elif ext == '.txt':
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    content = clean_text(text)
                    word_count = len(content)
                else:
                    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    content = clean_text(text)
                    word_count = len(content)
                
                # 清理临时文件
                shutil.rmtree(tempdir, ignore_errors=True)
                
                if word_count < 100:
                    raise FileProcessingError("AZW3内容过少")
                
                logger.info(f"AZW3读取成功，字数: {word_count:,}")
                return content, word_count
                
            finally:
                shutil.rmtree(tempdir, ignore_errors=True)
                
        except ImportError:
            # 回退到MOBI解析
            logger.warning("mobi库未安装，回退到基础解析")
            return self._read_mobi(file_path)
        except Exception as e:
            raise FileProcessingError(f"AZW3读取失败: {str(e)}。建议安装mobi库或转换为EPUB格式")
    
    def _read_html(self, file_path: str) -> Tuple[str, int]:
        """读取HTML文件"""
        import html
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        # 移除script和style标签
        html_content = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        # 移除HTML标签
        html_content = re.sub(r'<[^>]+>', '', html_content)
        # 解码HTML实体
        text_content = html.unescape(html_content)
        
        cleaned = clean_text(text_content)
        return cleaned, len(cleaned)
    
    def _read_docx(self, file_path: str) -> Tuple[str, int]:
        """读取DOCX文件"""
        logger.info(f"读取DOCX文件: {os.path.basename(file_path)}")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 读取段落
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # 读取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            text_parts.append(text)
            
            if not text_parts:
                raise FileProcessingError("DOCX文件无文本内容")
            
            full_content = '\n'.join(text_parts)
            cleaned = clean_text(full_content)
            
            logger.info(f"DOCX读取成功，字数: {len(cleaned):,}")
            return cleaned, len(cleaned)
            
        except Exception as e:
            raise FileProcessingError(f"DOCX读取失败: {str(e)}")
    
    def _read_doc(self, file_path: str) -> Tuple[str, int]:
        """读取DOC文件，优先用LibreOffice转换，备用antiword"""
        logger.info(f"读取DOC文件: {os.path.basename(file_path)}")
        
        # 方法1：LibreOffice转换
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # 转换为docx
                result = subprocess.run([
                    "soffice", "--headless", "--convert-to", "docx",
                    file_path, "--outdir", temp_dir
                ], capture_output=True, timeout=30)
                
                if result.returncode == 0:
                    # 查找转换后的文件
                    base_name = Path(file_path).stem
                    docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                    
                    if os.path.exists(docx_path):
                        return self._read_docx(docx_path)
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        # 方法2：antiword
        try:
            result = subprocess.run(
                ["antiword", file_path], 
                capture_output=True, 
                text=True, 
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout.strip():
                cleaned = clean_text(result.stdout)
                logger.info(f"DOC读取成功，字数: {len(cleaned):,}")
                return cleaned, len(cleaned)
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        
        raise FileProcessingError(
            "DOC读取失败。请安装LibreOffice或antiword，或将文件转换为DOCX/PDF格式"
        )
    

# 保持向后兼容的函数接口
def read_document(file_path: str) -> Tuple[str, int]:
    """
    向后兼容的文档读取函数
    
    Args:
        file_path: 文档路径
        
    Returns:
        Tuple[str, int]: (文本内容, 字符数)
    """
    reader = DocumentReader()
    return reader.read(file_path)


